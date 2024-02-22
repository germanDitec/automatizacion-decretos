from docx import Document
from docx.shared import Pt
from docx.shared import Inches


def add_table_to_word_document(doc, table):
    if table is None:
        return

    max_cols = max(len(row) for row in table)

    doc_table = doc.add_table(
        rows=len(table), cols=max_cols, style='Table Grid')

    for col in doc_table.columns:
        col.width = Pt(100)

    for i, row in enumerate(table):
        for j, cell in enumerate(row):
            doc_table.cell(i, j).text = cell

    for j in range(max_cols):
        doc_table.cell(0, j).paragraphs[0].runs[0].bold = True

