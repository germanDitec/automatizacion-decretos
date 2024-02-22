from flask import (Blueprint, render_template, request, flash, send_file,
                   redirect, url_for, Response, g, send_from_directory, session, current_app, abort)
from werkzeug.utils import secure_filename
from werkzeug.exceptions import HTTPException
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import concurrent.futures
import shortuuid
import os
import json
import mammoth
from app.auth import login_required
from app.db import get_db

from app.data_handling.data_extracting import (extract_table_from_pdf, extract_num_decreto, 
                                               extract_paragraphs_containing_keyword, extract_date_from_keyword, extract_date_from_last_page, extract_last_page, obtener_direccion, obtener_propuesta, extract_page_containing_keyword)

from app.data_handling.data_managing import (get_rechazados_text, add_decretos_lineas,
                                             add_decretos, formate_date_text, informe_to_sharepoint, decreto_to_sharepoint, add_visto)

from app.data_handling.data_writing import add_table_to_word_document


from openai import OpenAI
import json


bp = Blueprint('home', __name__, url_prefix="/")

SESSION_TYPE = 'filesystem'
ALLOWED_EXTENSIONS = {'pdf'}
# CUENTA SHAREPOINT, DEBE IR EN VARIABLES DE SESIÓN
server_url = "https://immaipu.sharepoint.com/"
site_url = server_url + "sites/AutomatizacionDecretos"

client = OpenAI(
    api_key=os.environ.get('OPENAI_API_KEY'),
)

image_path = os.path.join(os.path.dirname(
    __file__), "media/encabezado_decretos.png")


def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@bp.route("/", methods=["GET", "POST"])
@login_required
def index():
    """
    Esta función es la encargada de manejar el formulario de generación de decretos.
    Además también transforma algunos datos con el fin de entregar los datos limpios
    a la función de generar el documento final.
    """
    if request.method == 'POST' and 'file' in request.files:
        idp = request.form.get('idp')
        titulo = request.form.get('titulo')
        cdp = request.form.get('cdp')
        datecdp = request.form.get('datecdp')
        cuenta = request.form.get('cuenta')
        datecompra = request.form.get('datecompra')
        direccion = request.form.get('direccion')
        concejo = request.form.get('concejo')
        acuerdo = request.form.get('acuerdo')
        sesion = request.form.get('sesion')
        datesesion = request.form.get('datesesion')
        secretaria = request.form.get('secretaria')

        propuesta = request.form.get('propuesta')
        tipo_compra = request.form.get('compra')

        if concejo != 'on':
            acuerdo = None
            sesion = None
            datesesion = None
            secretaria = None
            concejo_bool = False
        else:
            concejo_bool = True

        error = None

        if not all([cdp, datecdp, datecompra, direccion]):
            error = "Por favor, rellena todos los campos."
            return render_template("home/index.html")

        valor_direccion = obtener_direccion(direccion)
        if direccion == "0":
            error = "Por favor, rellena todos los campos."

        if propuesta == "0":
            error = "Por favor, rellena todos los campos."

        pdf_file = request.files['file']
        if pdf_file and allowed_file(pdf_file.filename):
            filename = secure_filename(pdf_file.filename)
            upload_folder = 'app/static/documents_folder'
            os.makedirs(upload_folder, exist_ok=True)
            file_path = os.path.join(upload_folder, filename)
            pdf_file.save(file_path)
            informe_to_sharepoint(server_url, current_app.config['MAIL_USERNAME'],
                                  current_app.config['MAIL_PASSWORD'], site_url, file_path, filename, valor_direccion)
            db, c = get_db()
            c.execute(
                """
                INSERT INTO informe(idp, propuesta_id, direccion_id, titulo, cdp, fecha_cdp, 
                fecha_compra, cuenta, concejo, acuerdo_concejo, numero_sesion, fecha_sesion, 
                secretaria, tipo_compra, created_by) 
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                SELECT SCOPE_IDENTITY() as informe_id
                """, (idp, propuesta, direccion, titulo, cdp, datecdp, datecompra, cuenta, concejo_bool, acuerdo, sesion, datesesion, secretaria, tipo_compra, g.user[0])
            )

            session['informe_id'] = c.fetchone()[0]

            db.commit()

            session['process_ejected'] = False
            return redirect(url_for('home.generate_word', filename=filename, idp=idp, cdp=cdp, datecdp=datecdp, datecompra=datecompra,
                                    direccion=direccion, concejo=concejo, acuerdo=acuerdo, sesion=sesion, datesesion=datesesion,
                                    secretaria=secretaria, cuenta=cuenta, propuesta=propuesta, tipo_compra=tipo_compra, titulo=titulo))
        else:
            error = "El documento debe tener un formato PDF"

        flash(error, 'error')
    return render_template("home/index.html")


def get_rechazados(pdf_path):
    """
    En esta función se realiza la petición a la api de GPT para obtener los rechazados
    se retorna un json y también el coste de los prompts
    """
    try:
        input_promt = """Verifica si hay proponentes que fuéron rechazados, y si es así retornalos en formato JSON y hazlo con la siguiente estructura:
            rechazados: [
            nombre: nombre empresa,
            RUT: rut empresa,
            motivo: motivo de rechazo ]
            Y el caso de que no haya ningún proponente rechazado, retorna este json vacio:
                rechazados:[]
            Solo quiero el JSON, no incluyas texto adicional, haz esto en base a este texto. {}"""

        paragraphs_propmt = input_promt.format(extract_page_containing_keyword(
            pdf_path, "OBSERVACIONES DEL ACTO DE APERTURA"))
        rechazados_prompt = client.chat.completions.create(
            model="gpt-4-1106-preview",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "Quiero que analices la información que te estoy pasando y quiero que me retornes lo que te pido, solo lo que te pido, sin añadir texto adicional. El contexto del texto es sobre una comisión evaluadora."},
                {"role": "user", "content": paragraphs_propmt}
            ]
        )
        data = rechazados_prompt.choices[0].message.content
        prompt_tokens = rechazados_prompt.usage.prompt_tokens
        completion_tokens = rechazados_prompt.usage.completion_tokens
        datos = json.loads(data)
        rechazados = datos.get("rechazados", [])
        return rechazados, prompt_tokens, completion_tokens

    except:
        return [], 0, 0


def get_inadmisibles(pdf_path):
    """
    En esta función se realiza la petición a la api de GPT para obtener los inadmisibles
    se retorna un json y también el coste de los prompts
    """
    try:
        input_promt = """Verifica si en el parrafo de 'ADMISIBILIDAD' hay algún proponente que es inadmisible, y si es así retornalos en formato JSON y hazlo con la siguiente estructura: 
            inadmisibles: [
            nombre: nombre empresa, 
            RUT: rut empresa
            ]

            Y el caso de que no haya ningún proponente rechazado, retorna este json vacio:
                inadmisibles:[] 
            Solo quiero el JSON, no incluyas texto adicional, haz esto en base a este texto. {}"""

        paragraphs_propmt = input_promt.format(extract_paragraphs_containing_keyword(
            pdf_path, "ADMISIBILIDAD"))
        rechazados_prompt = client.chat.completions.create(
            model="gpt-4-1106-preview",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "Quiero que analices la información que te estoy pasando y quiero que me retornes lo que te pido, solo lo que te pido, sin añadir texto adicional. El contexto del texto es sobre una comisión evaluadora."},
                {"role": "user", "content": paragraphs_propmt}
            ]
        )
        data = rechazados_prompt.choices[0].message.content
        prompt_tokens = rechazados_prompt.usage.prompt_tokens
        completion_tokens = rechazados_prompt.usage.completion_tokens
        datos = json.loads(data)
        inadmisibles = datos.get("inadmisibles", [])
        print("inadmisibles:", inadmisibles)
        return inadmisibles, prompt_tokens, completion_tokens

    except:
        return [], 0, 0


def get_evaluacion(pdf_path):
    """
    En esta función se realiza la petición a la api de GPT para obtener la evaluación
    utilizando una función que extrae páginas a través de una palabra clave
    se retorna un json y también el coste de los prompts
    """
    try:
        evaluacion_input = "Quiero que hagas un resumen de la evaluación del o los proponentes desde el parrafo de 'ADMISIBILIDAD' hasta antes del párrafo de 'CONCLUSIÓN', redacta el resumen sin especificar que es lo que te pedí. También quiero que me listes con viñetas el resumen y si hay proponentes que fueron declarados inadmisibles detallalos, recuerda retornar en texto sin formato, haz esto en base a esta información: {}"
        evaluacion_page = evaluacion_input.format(
            extract_paragraphs_containing_keyword(pdf_path, "ADMISIBILIDAD"))

        evaluacion_prompt = client.chat.completions.create(
            model="gpt-4-1106-preview",
            messages=[
                {"role": "system", "content": "Quiero que me ayudes analizando el texto que te proporciono y me retornes lo que te pido, el contexto del texto es sobre una comisión evaluadora que ha realizado un informe sobre los proponentes para adjudicar y quiero que hagas un resumen de la evaluación. Esto ya pasó entonces escribe en tiempo pasado, no quiero que agregues porcentajes y utiliza un lenguaje muy formal, además no quiero que incluyas lo que concluyó la comisión para adjudicar"},
                {"role": "user", "content": evaluacion_page}
            ]
        )

        data_evaluacion = evaluacion_prompt.choices[0].message.content
        prompt_tokens = evaluacion_prompt.usage.prompt_tokens
        completion_tokens = evaluacion_prompt.usage.completion_tokens
        return data_evaluacion, prompt_tokens, completion_tokens

    except Exception as e:
        return f"No se pudo realizar la evaluación: {e}", 0, 0


def get_datos_adjudicacion(pdf_path):
    """
    En esta función se realiza la petición a la api de GPT para obtener los datos de los proponentes a adjudicación
    se retorna un json y también el coste de los prompts
    """
    try:
        adjudicacion = """Quiero que me retornes nombre, RUT y monto total de la o las empresas que hayan sido nombradas. Esto en formato JSON, extrae la información desde el párrafo de 'PROPOSICIÓN DE ADJUDICACIÓN'. Retorna solo el JSON con la siguiente estructura.
        En el caso de que existan más lineas en el parrafo de 'PROPOSICIÓN DE ADJUDICACION', retorna con el siguiente formato:
            empresas: [
            nombre: nombre empresa,
            RUT: rut empresa,
            linea: linea,
            total: total,
            plazo: plazo
            ]
            Y en el caso de que exista solamente una empresa retorna con el siguiente formato:
            empresas: [
            nombre: nombre empresa,
            RUT: rut empresa,
            total: total,
            plazo: plazo
            ]
            Esto en base a todo este texto, toma en cuenta solamente el parrafo de la 'PROPOSICIÓN DE ADJUDICACION', lo demás ignoralo. {}
            """
        adjudicacion_page = adjudicacion.format(extract_last_page(pdf_path))
        adjudicacion_prompt = client.chat.completions.create(
            model="gpt-4-1106-preview",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "Quiero que analices el texto que te estoy entregando y me retornes un JSON, solo con lo que te pido, sin añadir texto adicional. El contexto del texto es sobre una comisión evaluadora."},
                {"role": "user", "content": adjudicacion_page}
            ]
        )

        data_adjudicacion = adjudicacion_prompt.choices[0].message.content
        prompt_tokens = adjudicacion_prompt.usage.prompt_tokens
        completion_tokens = adjudicacion_prompt.usage.completion_tokens
        data = json.loads(data_adjudicacion)
        empresas = data.get("empresas", [])
        return empresas, prompt_tokens, completion_tokens

    except Exception as e:
        return [], 0, 0


def get_desiertas(pdf_path):
    """
    En esta función se realiza la petición a la api de GPT para obtener las lineas desiertas
    se retorna un json y también el coste de los prompts
    """
    try:
        input_promt = """{}
        Analiza el párrafo de 'PROPOSICIÓN DE ADJUDICACION' y retorna un JSON que contenga las líneas desiertas, representadas por el formato: 
            desiertas:[
                    linea: numero linea o números de linea,
                    motivo: la razon de porque está desierta
                ]
        donde se listarán los números de línea que no tienen ofertas o adjudicaciones.
        Y el caso de que no haya ninguna linea desierta, retorna este json vacio: 
            desiertas:[]
        Solo quiero el JSON, no incluyas texto adicional, haz esto en base al texto que te estoy proporcionando"""

        desiertas_prompt = input_promt.format(
            extract_last_page(pdf_path))

        desiertas = client.chat.completions.create(
            model="gpt-4-1106-preview",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "Quiero que analices la información que te estoy pasando y quiero que me retornes la información en formato JSON."},
                {"role": "user", "content": desiertas_prompt}
            ]
        )
        data = desiertas.choices[0].message.content
        prompt_tokens = desiertas.usage.prompt_tokens
        completion_tokens = desiertas.usage.completion_tokens
        datos = json.loads(data)
        desiertas = datos.get("desiertas", [])
        return desiertas, prompt_tokens, completion_tokens

    except:
        return [], 0, 0


# Lista de funciones para llamar luego en la función de generar el decreto
functions_list = [get_rechazados, get_inadmisibles,
                  get_evaluacion, get_datos_adjudicacion, get_desiertas]


@bp.route("/generate/<filename>", methods=['GET'])
@login_required
def generate_word(filename):
    """
    En esta función se encarga de generar el documento final.
    Recibe la información del formulario a través de parametros y luego se utilizan 
    para generar el documento docx.
    """
    if filename.endswith(".pdf"):
        try:
            regenerate = request.args.get('regenerate', 'false')
            if regenerate.lower() == 'true':
                session['process_ejected'] = False

            doc = Document()
            pdf_path = os.path.join('app/static/documents_folder', filename)
            cdp = request.args.get('cdp')
            cuenta = request.args.get('cuenta')
            concejo = request.args.get("concejo")

            acuerdo = request.args.get('acuerdo')
            sesion = request.args.get('sesion')
            datesesion = request.args.get('datesesion')
            if datesesion is not None:
                datesesion_str = formate_date_text(datesesion)
            else:
                datesesion_str = None
            secretaria = request.args.get('secretaria')
            datecdp = request.args.get('datecdp')
            datecompra = request.args.get('datecompra')
            datecdp_str = formate_date_text(datecdp)
            datecompra_str = formate_date_text(datecompra)
            direccion = request.args.get('direccion')
            valor_direccion = obtener_direccion(direccion)
            n_decreto = extract_num_decreto(pdf_path)
            idp = request.args.get('idp')
            titulo = request.args.get('titulo')
            titulo = titulo.upper()
            propuesta = request.args.get('propuesta')
            valor_propuesta = obtener_propuesta(propuesta)
            tipo_compra = request.args.get('tipo_compra')
            fecha_apertura = extract_date_from_keyword(
                pdf_path, "El Acto de Apertura")
            fecha_informe = extract_date_from_last_page(pdf_path)
            fecha_decreto = extract_date_from_keyword(
                pdf_path, "Decreto Alcaldicio")

            considerando_primer = f"1.- Que, con fecha {datecompra_str}, se publicó en el Sistema de Información de Compras y Contratación Pública (www.mercadopublico.cl), la propuesta pública {idp}, denominada {titulo}, según Bases administrativas y Técnicas, aprobadas por Decreto Alcaldicio N°{n_decreto} de fecha {fecha_decreto}"
            considerando_segundo = "2.- Que, en el llamado a licitación los siguientes proponentes presentaron ofertas:"
            considerando_tercer = ""
            considerando_cuarto = ""
            considerando_quinto = ""
            considerando_sexto = ""
            considerando_septimo = None

            if session.get('process_ejected'):
                return render_template('home/process.html', html=session.get('generated_document', ''), filename=filename,
                                       cdp=cdp, datecdp=datecdp, datecompra=datecompra, cuenta=cuenta,
                                       concejo=concejo, acuerdo=acuerdo, sesion=sesion,
                                       datesesion=datesesion, secretaria=secretaria,
                                       direccion=direccion, tipo_compra=tipo_compra,
                                       propuesta=propuesta, idp=idp, titulo=titulo)

            # CONCURRENCY
            with concurrent.futures.ThreadPoolExecutor() as executor:
                results = list(executor.map(
                    lambda func: func(pdf_path), functions_list))

            rechazados, prompt_rechazados, completion_rechazados = results[0]
            inadmisibles, prompt_inadmisibles, completion_inadmisibles = results[1]
            data_evaluacion, prompt_evaluacion, completion_evaluacion = results[2]
            empresas_adjudicadas, prompt_adjudicacion, completion_adjudicacion = results[3]
            desiertas, prompt_desiertas, completion_desiertas = results[4]

            total_prompt_tokens = prompt_rechazados + prompt_inadmisibles + \
                prompt_evaluacion + prompt_adjudicacion + prompt_desiertas
            total_completion_tokens = completion_rechazados + completion_inadmisibles + \
                completion_evaluacion + completion_adjudicacion + completion_desiertas

            price_prompt = (total_prompt_tokens / 1000) * 0.01
            price_completion = (total_completion_tokens / 1000) * 0.03
            total_price = round(price_prompt + price_completion, 2)

            fragmento_adjudicadas = []
            posee_linea = False
            for empresa in empresas_adjudicadas:
                nombre_adjudicada = empresa.get("nombre", "")
                rut_adjudicada = empresa.get("RUT", "")
                linea = empresa.get("linea", None)
                total = empresa.get("total", "")
                fragmento_empresas = f"{nombre_adjudicada} RUT {rut_adjudicada}, para la línea {linea}"
                fragmento_adjudicadas.append(fragmento_empresas)
                proponentes_adjudicados = "; ".join(fragmento_adjudicadas)
                if linea is not None:
                    lista_empresa_adjudicada = f"- Del resultado de la evaluación, se establece que la oferta presentada por los proponentes {proponentes_adjudicados} . Se propone sean adjudicadas dado que son las ofertas más convenientes para los intereses municipales."
                    posee_linea = True
                else:
                    lista_empresa_adjudicada = f"- Del resultado de la evaluación, se establece que la oferta presentada por el proponente {nombre_adjudicada}, RUT {rut_adjudicada}. Se propone sea adjudicada dado que es la oferta más conveniente para los intereses muncipales."

            if rechazados:
                considerando_tercer = f"\n3.- Que, en el Acta de Apertura de las ofertas, de fecha {fecha_apertura}, se rechazaron las siguientes ofertas:"
                proponentes_rechazados, motivo_rechazo = get_rechazados_text(
                    rechazados)
                lista_rechazados = f"La de los proponentes {proponentes_rechazados}, por motivo de {motivo_rechazo}."

            else:
                considerando_tercer = "\n3.- Que, en el Acto de Apertura de las ofertas, de fecha {}, no existen ofertas rechazadas.".format(
                    fecha_apertura)
                lista_rechazados = ""

            considerando_cuarto = f"\n4.- Que, de acuerdo con el informe de Evaluación de Ofertas, de fecha {fecha_informe}, la comisión evaluadora estableció lo siguiente:"

            # En el caso de que el usuario haya seleccionado que la licitación sí pasó por concejo se agrega el texto correspondiente
            if concejo == 'on':
                considerando_quinto = f"\n5.- Que, la propuesta de adjudicación cuenta con el Acuerdo N° {acuerdo}, adoptado en Sesión Ordinaria N° {sesion} de fecha {datesesion_str}, según consta en Certificado N° {secretaria} de Secretaria Municipal, del Honorable Concejo Municipal."
                considerando_sexto = f"\n6.- Que, se cuenta con la disponibilidad presupuestaria para este fin, según da cuenta el Certificado de Factibilidad N° {cdp}, de fecha {datecdp_str}."
                considerando_septimo = f"\n7.- Que, en el Numeral 13 de las Bases Administrativas, establece que la Unidad Técnica responsable de supervisar la ejecución de {valor_propuesta.upper()} será la {valor_direccion}."
            # En el caso de que el usuario no haya seleccionado que la licitación pasó por concejo se agrega el texto correspondiente
            else:
                considerando_quinto = f"\n5.- Que, se cuenta con la disponibilidad presupuestaria para este fin, según da cuenta el Certificado de Factibilidad N° {cdp}, de fecha {datecdp_str}."
                considerando_sexto = f"\n6.- Que, en el Numeral 13 de las Bases Administrativas, establece que la Unidad Técnica responsable de supervisar la ejecución de {valor_propuesta.upper()} será la {valor_direccion}."

            table = extract_table_from_pdf(pdf_path)
            output_file_path = pdf_path.replace(
                ".pdf", "-PLANTILLA_DECRETO.docx")
            header = doc.sections[0].header
            paragraph = header.paragraphs[0]
            logo_run = paragraph.add_run()
            # Imagen de pie de página del documento
            logo_run.add_picture(
                image_path, width=Inches(7), height=Inches(0.5))
            fecha_decreto_alcaldicio = doc.add_paragraph()
            fecha_decreto_alcaldicio_r = fecha_decreto_alcaldicio.add_run(
                "Maipú,"
            )
            fecha_decreto_alcaldicio_r.jc = 'left'
            fecha_decreto_alcaldicio_r.font.size = Pt(10)
            fecha_decreto_alcaldicio_r.bold = True
            titulo_decreto = doc.add_paragraph()
            titulo_decreto_r = titulo_decreto.add_run("DECRETO ALCALDICIO N° ")
            titulo_decreto_r.jc = 'left'
            titulo_decreto_r.font.size = Pt(10)
            titulo_decreto_r.bold = True
            visto_p = doc.add_paragraph()
            visto_r = visto_p.add_run("VISTO:")
            visto_r.bold = True
            add_visto(doc, n_decreto, fecha_decreto, acuerdo, sesion, datesesion_str, secretaria,
                      idp, titulo, fecha_apertura, fecha_informe, cdp, datecdp_str, concejo)
            considerando_p = doc.add_paragraph()
            considerando_r = considerando_p.add_run("CONSIDERANDO:")
            considerando_r.bold = True
            doc.add_paragraph(considerando_primer)
            doc.add_paragraph(considerando_segundo)
            add_table_to_word_document(doc, table)
            doc.add_paragraph(considerando_tercer)
            doc.add_paragraph(lista_rechazados, style="List Bullet")
            doc.add_paragraph(considerando_cuarto)
            doc.add_paragraph(data_evaluacion)
            doc.add_paragraph(lista_empresa_adjudicada)
            doc.add_paragraph(considerando_quinto)
            doc.add_paragraph(considerando_sexto)
            if considerando_septimo is not None:
                doc.add_paragraph(considerando_septimo)
            decreto_p = doc.add_paragraph()
            decreto_r = decreto_p.add_run("DECRETO:")

            decreto_r.bold = True
            if posee_linea:
                add_decretos_lineas(doc, idp, titulo, rechazados, inadmisibles,
                                    empresas_adjudicadas, valor_direccion, cuenta, tipo_compra, propuesta, valor_propuesta, desiertas)
            else:
                add_decretos(doc, idp, titulo, rechazados, inadmisibles, valor_direccion, cuenta,
                             tipo_compra, propuesta, valor_propuesta, desiertas, nombre_adjudicada, total, rut_adjudicada)

            doc.save(output_file_path)

            db, c = get_db()

            c.execute(
                """
                INSERT INTO decreto(informe_id, costo_total, propuesta_id, direccion_id, created_by)
                VALUES (%s, %s, %s, %s, %s)
                """,
                (int(session['informe_id']), total_price,
                 propuesta, direccion, session['user_id'])
            )

            db.commit()

            with open(output_file_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                # El resultado se transforma en HTML para mostrar el docx en la página de resultado
                html = result.value
                session['generated_document'] = html
                # Luego de realizar todo el proceso, se indica que ya se ha ejectutado el proceso
                session['process_ejected'] = True

            output_filename = filename.replace(
                ".pdf", f"-PLANTILLA_DECRETO.docx")

            # Se envia el decreto al Sharepoint
            decreto_to_sharepoint(
                server_url, current_app.config['MAIL_USERNAME'], current_app.config['MAIL_PASSWORD'], site_url, output_file_path, output_filename, valor_direccion)

            # Se carga la plantilla cuando se procesa todo, con todos los argumentos que se utilizan en procces.html
            return render_template('home/process.html', html=html, filename=filename,
                                   cdp=cdp, datecdp=datecdp, datecompra=datecompra,
                                   cuenta=cuenta, concejo=concejo, acuerdo=acuerdo,
                                   sesion=sesion, datesesion=datesesion, secretaria=secretaria,
                                   direccion=direccion, tipo_compra=tipo_compra, propuesta=propuesta, idp=idp, titulo=titulo)
        except HTTPException as e:
            if e.code == 400:
                return render_template('errors/400.html'), 400

        except Exception as e:
            return render_template('errors/error_general.html', error=str(e)), 500


# Endpoint para descargar la plantilla del decreto
@bp.route("/download/<filename>")
def download_file(filename):
    output_file_path = filename.replace(".pdf", f"-PLANTILLA_DECRETO.docx")
    return send_from_directory('static/documents_folder', output_file_path, as_attachment=True)


@ bp.errorhandler(400)
def bad_request(e):
    return render_template('errors/400.html'), 400
